from docx import Document


paths = [
    r"C:/Users/cps/Downloads/Report7_Final Project Report (1).docx",
    r"C:/Users/cps/Downloads/FoodResQ_Report5_Test_Documentation.docx",
]

for path in paths:
    doc = Document(path)
    print("FILE", path)
    print("paras", len(doc.paragraphs), "tables", len(doc.tables))
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if path.endswith("(1).docx"):
            if 115 <= index <= 170 or text.startswith("V.") or text.startswith("VI."):
                print(index, paragraph.style.name, text[:240])
        else:
            if index < 170:
                print(index, paragraph.style.name, text[:240])
    print()
