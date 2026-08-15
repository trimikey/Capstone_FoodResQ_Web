from docx import Document


path = r"D:/Capstone_FoodResQ_Web/outputs/report7-final/FoodResQ_Report7_Final_Project_Report_testing_filled.docx"
doc = Document(path)
print("paras", len(doc.paragraphs), "tables", len(doc.tables))

for index, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    if not text:
        continue
    if text.startswith("V.") or text.startswith("VI.") or text in {
        "1. Scope of Testing",
        "2. Test Strategy",
        "2.1 Testing Types",
        "2.2 Test Levels",
        "2.3 Supporting Tools",
        "3. Test Plan",
        "3.1 Human Resources",
        "3.2 Test Environment",
        "3.3 Test Milestones",
        "4. Test Cases",
        "5. Test Reports",
    }:
        print(index, paragraph.style.name, text)

print("statistics table sample")
for row in doc.tables[-2].rows[-5:]:
    print([cell.text for cell in row.cells])

all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
bad = [term for term in ["BookSwapHub", "Course payment", "Learner", "Coach"] if term in all_text]
print("bad_terms", bad)
