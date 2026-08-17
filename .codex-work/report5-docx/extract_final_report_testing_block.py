from docx import Document


path = r"C:/Users/cps/Downloads/Report7_Final Project Report.docx"
doc = Document(path)

start = 120
end = 160
for index in range(start, min(end, len(doc.paragraphs))):
    paragraph = doc.paragraphs[index]
    text = paragraph.text.strip()
    if text:
        print(index, paragraph.style.name, text)
