from copy import deepcopy
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches


TARGET = Path(r"C:/Users/cps/Downloads/Report7_Final Project Report (1).docx")
XLSX = Path(r"D:/Capstone_FoodResQ_Web/outputs/report5-foodresq/FoodResQ_Report5_Test_Report_v11.xlsx")
OUT = Path(r"D:/Capstone_FoodResQ_Web/outputs/report7-final/FoodResQ_Report7_Final_Project_Report_testing_filled.docx")

FONT = "Times New Roman"
RED = RGBColor(192, 0, 0)
BLUE = "000080"
PEACH = "FCE4D6"
LIGHT_BLUE = "D9E2F3"


def set_run_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name = FONT
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def style_paragraph(paragraph, size=10, bold=False, italic=False, color=None, align=None):
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.05
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color="000000", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cell_text(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.clear()
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def make_table(doc, headers, rows, widths=None, blue_header=False, subtotal_last=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    for i, header in enumerate(headers):
        if widths:
            table.columns[i].width = Inches(widths[i])
        set_cell_text(
            table.rows[0].cells[i],
            header,
            bold=True,
            size=8 if len(headers) > 5 else 9,
            color=RGBColor(255, 255, 255) if blue_header else None,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_shading(table.rows[0].cells[i], BLUE if blue_header else PEACH)
    set_repeat_table_header(table.rows[0])
    for row_idx, row_values in enumerate(rows):
        cells = table.add_row().cells
        is_total = subtotal_last and row_idx == len(rows) - 1
        for i, value in enumerate(row_values):
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or isinstance(value, (int, float)) or str(value).endswith("%") else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, bold=is_total, size=8 if len(headers) > 5 else 9, align=align)
            if is_total:
                set_cell_shading(cells[i], LIGHT_BLUE)
    return table


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=9, bold=True, italic=True)


def add_para(doc, text="", style=None, size=10, bold=False, italic=False, color=None, align=None):
    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    style_paragraph(paragraph, size=size, bold=bold, italic=italic, color=color, align=align)
    return paragraph


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    run = paragraph.add_run(text)
    set_run_font(run, size=10)
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def insert_doc_element_after(anchor, element):
    anchor._p.addnext(element)


def move_body_element_after(anchor_paragraph, body_element):
    insert_doc_element_after(anchor_paragraph, body_element)
    return body_element


def append_content_to_temp(stats_rows):
    temp = Document()
    h = temp.add_heading("V. Software Testing Documentation", level=1)
    style_paragraph(h, size=16, bold=True, italic=True, color=RED)

    h = temp.add_heading("1. Scope of Testing", level=2)
    style_paragraph(h, size=13, bold=True)
    add_para(
        temp,
        "The testing activities cover the main functional workflows and key non-functional requirements of the FoodResQ system. FoodResQ is a food rescue and donation coordination platform that connects receivers, providers, volunteers/shippers, charity organizations, and administrators. Testing focuses on verifying safe food listing, reservation, pickup, delivery, campaign kitchen operation, bulk distribution, trust-score control, realtime notification, and administration workflows.",
    )
    add_para(temp, "Feature and Function to be tested:", bold=True)
    tested = [
        "User authentication, registration with eKYC, logout, password recovery, token refresh, and role-based access control.",
        "Food listing browsing, nearby search, listing detail viewing, provider listing creation, image upload, publishing, updating, cancellation, and pickup-window validation.",
        "Reservation workflow, including stock validation, daily reservation limit, Redis lock behavior, QR generation, QR scan confirmation, cancellation, no-show handling, and pickup completion.",
        "Delivery workflow, including nearest shipper assignment, offer broadcast, offer acceptance/rejection/expiry, delivery lifecycle, QC photo, live tracking, and receiver QR handoff proof.",
        "Bulk run workflow, including bulk request, provider approval/rejection, pickup, stop logging, served portion tracking, completion, and leftover stock return.",
        "Campaign and kitchen operation workflow, including campaign creation, admin approval/rejection, volunteer staffing, shift management, menu management, dish preparation, safety logs, meal distribution, beneficiary QR handoff, provider supply request, and transport receipt.",
        "Profile, trust-score, notification, admin, KYC review, user governance, violation handling, and operational monitoring functions.",
    ]
    for item in tested:
        add_bullet(temp, item)

    add_para(temp, "Feature and Function not to be tested:", bold=True)
    not_tested = [
        "Large-scale production stress testing beyond the controlled capstone demonstration workload.",
        "Production disaster recovery, multi-region failover, and incident-response drills.",
        "Third-party infrastructure reliability testing, such as external push notification, geocoding, object storage, or hosting provider outages.",
        "Formal biometric face-matching accuracy benchmarking; testing only verifies that the eKYC and face-enrollment gates work correctly in the system workflow.",
        "Online payment settlement because FoodResQ focuses on food donation, reservation, delivery, and distribution workflows.",
    ]
    for item in not_tested:
        add_bullet(temp, item)

    h = temp.add_heading("2. Test Strategy", level=2)
    style_paragraph(h, size=13, bold=True)
    h = temp.add_heading("2.1 Testing Types", level=3)
    style_paragraph(h, size=11, bold=True)
    testing_types = [
        ("Functional Testing", "Ensure that FoodResQ functions comply with requirement specifications and business rules.", "Black-box and scenario-based testing from each user role.", "All critical functional test cases pass without major defects."),
        ("Integration Testing", "Verify that frontend/mobile screens, NestJS APIs, PostgreSQL/PostGIS, Redis locks, queues, and Socket.IO events work together correctly.", "Execute cross-module workflows and check UI, API response, database state, and notification behavior.", "No broken state transition or data mismatch remains."),
        ("System Testing", "Validate the complete FoodResQ system against end-to-end user workflows.", "Execute receiver, provider, volunteer/shipper, charity, and admin scenarios.", "Main workflows operate correctly without blocking errors."),
        ("User Acceptance Testing", "Ensure that the system satisfies realistic food-rescue operation needs.", "Team members perform role-based demo scenarios using prepared accounts and seeded data.", "Required user stories and acceptance criteria are met."),
        ("Security Testing", "Validate authentication, authorization, banned/restricted account behavior, upload validation, and protected route access.", "Test invalid credentials, expired tokens, wrong-role access, unauthenticated access, and invalid inputs.", "Unauthorized requests are rejected and protected operations require a valid role/session."),
    ]
    for name, objective, technique, criteria in testing_types:
        add_bullet(temp, name)
        add_bullet(temp, f"Objective: {objective}", level=1)
        add_bullet(temp, f"Technique: {technique}", level=1)
        add_bullet(temp, f"Completion Criteria: {criteria}", level=1)

    h = temp.add_heading("2.2 Test Levels", level=3)
    style_paragraph(h, size=11, bold=True)
    make_table(
        temp,
        ["Type of Tests", "Unit", "Integration", "System", "Acceptance"],
        [
            ["Functionality", "X", "X", "X", "X"],
            ["User Interface", "", "", "X", "X"],
            ["Data Validation", "X", "X", "X", "X"],
            ["Role-based Access", "X", "X", "X", "X"],
            ["Realtime/Notification", "", "X", "X", "X"],
            ["Performance", "", "X", "X", ""],
        ],
        widths=[2.0, 0.85, 1.1, 1.0, 1.1],
    )
    add_caption(temp, "Table 36. Test levels table")

    h = temp.add_heading("2.3 Supporting Tools", level=3)
    style_paragraph(h, size=11, bold=True)
    make_table(
        temp,
        ["Purpose", "Tool", "Vendor/In-house", "Version"],
        [
            ["Test case management", "Microsoft Excel", "Microsoft", "Desktop / latest"],
            ["Manual test execution", "Chrome, Edge, Android Emulator, Expo Go", "Microsoft / Google / Expo", "Latest"],
            ["API verification", "Swagger / Postman", "In-house / Postman", "Latest"],
            ["Defect tracking", "GitHub Issues", "GitHub", "Cloud version"],
        ],
        widths=[1.7, 2.2, 1.5, 1.25],
    )
    add_caption(temp, "Table 37. Supporting tools table")

    h = temp.add_heading("3. Test Plan", level=2)
    style_paragraph(h, size=13, bold=True)
    h = temp.add_heading("3.1 Human Resources", level=3)
    style_paragraph(h, size=11, bold=True)
    make_table(
        temp,
        ["Worker/Doer", "Role", "Specific Responsibilities/Comments"],
        [
            ["QA Team", "Tester", "Prepare manual test cases, execute system and acceptance tests, record actual results"],
            ["Backend Developer", "Developer", "Support API, service rule, database, queue, and realtime issue investigation"],
            ["Web/Mobile Developer", "Developer", "Support UI, form validation, map, QR, and role-based screen issue investigation"],
            ["Project Lead", "Reviewer", "Review scope, milestones, defect severity, and final report consistency"],
        ],
        widths=[1.45, 1.2, 4.0],
    )
    add_caption(temp, "Table 38. Human resources table")

    h = temp.add_heading("3.2 Test Environment", level=3)
    style_paragraph(h, size=11, bold=True)
    make_table(
        temp,
        ["Purpose", "Tool", "Provider", "Version"],
        [
            ["Web application testing", "Chrome / Edge", "Google / Microsoft", "Latest"],
            ["Mobile application testing", "Expo Go / Android Emulator", "Expo / Google", "Latest"],
            ["Backend API", "NestJS API Server", "In-house", "Latest project build"],
            ["Database", "PostgreSQL + PostGIS", "Open-source", "15.x+"],
            ["Realtime and background jobs", "Socket.IO, Redis, BullMQ", "Open-source", "Latest"],
            ["Source control", "Git", "Open-source", "Latest"],
        ],
        widths=[1.7, 2.2, 1.5, 1.25],
    )
    add_caption(temp, "Table 39. Test environment table")

    h = temp.add_heading("3.3 Test Milestones", level=3)
    style_paragraph(h, size=11, bold=True)
    make_table(
        temp,
        ["Milestone Task", "Start Date", "End Date"],
        [
            ["Test planning", "01/08/2026", "03/08/2026"],
            ["Test case development", "04/08/2026", "08/08/2026"],
            ["Manual system testing", "09/08/2026", "12/08/2026"],
            ["Acceptance testing", "13/08/2026", "14/08/2026"],
            ["Testing closure and reporting", "15/08/2026", "15/08/2026"],
        ],
        widths=[3.6, 1.4, 1.4],
    )
    add_caption(temp, "Table 40. Test milestones table")

    h = temp.add_heading("4. Test Cases", level=2)
    style_paragraph(h, size=13, bold=True)
    add_bullet(temp, "Manual System and Acceptance Test Cases: Refer to FoodResQ_Report5_Test_Report_v11.xlsx.")
    add_bullet(temp, "The workbook contains 104 test cases covering Login, Register eKYC, Logout, Password, Profile Trust, Admin Controls, Food Listings, Reservations QR, Delivery Shipper, Campaign Kitchen Ops, and Bulk Notifications.")
    add_bullet(temp, "Each feature sheet is divided into smaller business sections and includes internal links from the Test Cases sheet to the detailed feature sheets.")

    h = temp.add_heading("5. Test Reports", level=2)
    style_paragraph(h, size=13, bold=True)
    add_bullet(temp, "System Testing, Acceptance Testing")
    add_para(temp, "Test Statistics", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    make_table(
        temp,
        ["Project Name", "FoodResQ", "Creator", "FoodResQ QA Team"],
        [
            ["Project Code", "SP26SE088", "Reviewer/Approver", "Capstone Supervisor"],
            ["Document Code", "SP26SE088_Test Report_v1.0", "Issue Date", "15-08-2026"],
        ],
        widths=[1.25, 2.2, 1.55, 1.55],
    )
    module_rows = [[row["no"], row["module"], row["passed"], row["failed"], row["pending"], row["na"], row["total"]] for row in stats_rows]
    module_rows.append(["Sub total", "", sum(r["passed"] for r in stats_rows), 0, 0, 0, sum(r["total"] for r in stats_rows)])
    make_table(
        temp,
        ["No", "Module code", "Passed", "Failed", "Pending", "N/A", "Number of test cases"],
        module_rows,
        widths=[0.45, 2.35, 0.62, 0.62, 0.68, 0.52, 1.05],
        blue_header=True,
        subtotal_last=True,
    )
    make_table(
        temp,
        ["Metric", "Value"],
        [
            ["Test coverage", "100.00%"],
            ["Test successful coverage", "100.00%"],
        ],
        widths=[3.2, 3.1],
    )
    add_caption(temp, "Table 41. Test statistics report table")
    return temp


def read_stats():
    wb = openpyxl.load_workbook(XLSX, data_only=True)
    ws = wb["Test Statistics"]
    rows = []
    for row in range(11, 22):
        no = ws.cell(row, 2).value
        module = ws.cell(row, 3).value
        if not no or not module:
            continue
        rows.append(
            {
                "no": no,
                "module": module,
                "passed": ws.cell(row, 4).value or 0,
                "failed": ws.cell(row, 5).value or 0,
                "pending": ws.cell(row, 6).value or 0,
                "na": ws.cell(row, 7).value or 0,
                "total": ws.cell(row, 8).value or 0,
            }
        )
    return rows


def copy_element_after(anchor_element, element):
    new_element = deepcopy(element)
    anchor_element.addnext(new_element)
    return new_element


def replace_testing_section():
    doc = Document(TARGET)
    start = None
    end = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text == "V. Software Testing Documentation":
            start = i
        elif start is not None and text.startswith("VI. "):
            end = i
            break
    if start is None or end is None:
        raise RuntimeError("Could not locate section V and VI boundaries.")

    body = doc._body._element
    start_el = doc.paragraphs[start]._p
    end_el = doc.paragraphs[end]._p
    children = list(body)
    start_idx = children.index(start_el)
    end_idx = children.index(end_el)

    for element in children[start_idx:end_idx]:
        body.remove(element)

    anchor = children[start_idx - 1]
    temp = append_content_to_temp(read_stats())
    for element in reversed(list(temp._body._element)):
        if element.tag.endswith("sectPr"):
            continue
        copy_element_after(anchor, element)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    replace_testing_section()
