import sys
from pathlib import Path
from docx import Document

src = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(r"C:/Users/cps/Downloads/Report5_Test Documentation.docx")
doc = Document(src)
out = Path(".codex-work/report5-docx/inspection.txt")
lines = []

lines.append(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")
for i, section in enumerate(doc.sections):
    lines.append(
        f"section {i}: page={section.page_width}x{section.page_height} margins="
        f"{section.left_margin},{section.right_margin},{section.top_margin},{section.bottom_margin}"
    )

lines.append("\nPARAGRAPHS")
for i, p in enumerate(doc.paragraphs):
    text = p.text.replace("\n", "\\n")
    if text.strip():
        lines.append(f"P{i:03d} style={p.style.name!r}: {text[:500]}")

lines.append("\nTABLES")
for ti, table in enumerate(doc.tables):
    lines.append(f"TABLE {ti}: rows={len(table.rows)} cols={len(table.columns)}")
    for ri, row in enumerate(table.rows[:30]):
        cells = []
        for cell in row.cells:
            text = cell.text.replace("\n", " | ").strip()
            cells.append(text[:220])
        lines.append(f"  R{ri:02d}: " + " || ".join(cells))
    if len(table.rows) > 30:
        lines.append("  ...")

out.write_text("\n".join(lines), encoding="utf-8")
print(out)
