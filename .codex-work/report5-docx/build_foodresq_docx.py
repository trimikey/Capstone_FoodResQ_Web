from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("outputs/report5-docx")
OUT = OUT_DIR / "FoodResQ_Report5_Test_Documentation_v2.docx"

FONT = "Times New Roman"
ACCENT_RED = RGBColor(192, 0, 0)
HEADER_BLUE = "1F2A7A"
PALE_PEACH = "FCE4D6"
PALE_BLUE = "D9E2F3"


MODULES = [
    ("AUTH", "Login", 6),
    ("REG", "Register eKYC", 7),
    ("LOGOUT", "Logout", 3),
    ("PWD", "Password", 4),
    ("PROFILE", "Profile Trust", 10),
    ("ADMIN", "Admin Controls", 8),
    ("LISTING", "Food Listings", 10),
    ("RES-QR", "Reservations QR", 9),
    ("DELIVERY", "Delivery Shipper", 14),
    ("CAMPAIGN", "Campaign Kitchen Ops", 16),
    ("BULK-NOTI", "Bulk Notifications", 5),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=10, color=None, align=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.clear()
    run = p.add_run(str(text))
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_table_borders(table, color="000000", size="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def make_table(doc, headers, rows, widths=None, header_fill=PALE_PEACH, blue_header=False, subtotal_row=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    if widths:
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=9, color=RGBColor(255, 255, 255) if blue_header else None, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], HEADER_BLUE if blue_header else header_fill)
    set_repeat_table_header(table.rows[0])
    for r_idx, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or str(value).isdigit() or str(value).endswith("%") else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, bold=subtotal_row and r_idx == len(rows), size=9, align=align)
            if subtotal_row and r_idx == len(rows):
                set_cell_shading(cells[i], PALE_BLUE)
    doc.add_paragraph()
    return table


def set_document_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10)

    for name, size, bold, color in [
        ("Heading 1", 16, True, ACCENT_RED),
        ("Heading 2", 13, True, RGBColor(0, 0, 0)),
        ("Heading 3", 11, True, RGBColor(0, 0, 0)),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10 if name != "Heading 1" else 14)
        style.paragraph_format.space_after = Pt(6)


def add_para(doc, text="", bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    if text:
        r = p.add_run(text)
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        r._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        r.font.size = Pt(10)
        r.bold = bold
        r.italic = italic
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    r._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    r.font.size = Pt(10)


def add_caption(doc, text):
    p = add_para(doc, text, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for run in p.runs:
        run.font.size = Pt(9)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    if level == 1:
        for run in p.runs:
            run.italic = True
    return p


def add_metadata_table(doc, title):
    p = add_para(doc, title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for run in p.runs:
        run.font.size = Pt(14)

    rows = [
        ["Project Name", "FoodResQ", "Creator", "QA Team"],
        ["Project Code", "SP26SE088", "Reviewer/Approver", "Project Supervisor"],
        ["Document Code", "SP26SE088_Test_Report_v1.0", "Issue Date", "15-08-2026"],
    ]
    table = make_table(doc, ["", "", "", ""], rows, widths=[1.25, 2.15, 1.55, 1.55])
    table._tbl.remove(table.rows[0]._tr)
    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        set_cell_text(row.cells[0], row.cells[0].text, bold=True, size=9)
        set_cell_text(row.cells[2], row.cells[2].text, bold=True, size=9)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    set_document_styles(doc)

    add_para(doc, "CAPSTONE PROJECT REPORT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    title = add_para(doc, "Report 5 - Software Test Documentation", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for run in title.runs:
        run.font.size = Pt(18)
    add_para(doc, "FoodResQ - Food Rescue and Donation Coordination System", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "- Ho Chi Minh City, August 2026 -", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    add_heading(doc, "I. Record of Changes", 1)
    add_para(doc, "*A - Added   M - Modified   D - Deleted", italic=True)
    changes = [
        ["15/08/2026", "M", "QA Team", "Reworked testing documentation to match the general report structure"],
        ["15/08/2026", "A", "QA Team", "Added FoodResQ scope, test strategy, test plan, and manual test report statistics"],
        ["15/08/2026", "M", "QA Team", "Aligned report content with FoodResQ workflows and the latest test case workbook"],
    ]
    make_table(doc, ["Date", "A* | M, D", "In charge", "Change Description"], changes, widths=[1.1, 1.0, 1.35, 3.65])

    add_heading(doc, "V. Testing Documentation", 1)

    add_heading(doc, "1. Scope of Testing", 2)
    add_para(
        doc,
        "The testing activities cover the main functional workflows and key non-functional requirements of the FoodResQ system, including user authentication, eKYC registration, food listing management, reservation and QR pickup, delivery coordination, bulk distribution, campaign/kitchen operation support, realtime notifications, trust-score handling, and administrative control functions.",
    )
    add_para(doc, "Feature and Function to be tested:", bold=True)
    for item in [
        "User authentication, registration with eKYC, logout, password recovery, token refresh, and role-based access.",
        "Food listing browsing, nearby search with map/location data, listing detail viewing, provider listing creation, publishing, update, cancellation, and image upload.",
        "Reservation creation, stock validation, daily-limit validation, pickup-window validation, QR generation, QR scan confirmation, cancellation, no-show handling, and completion.",
        "Delivery flow, including shipper availability, nearest-shipper offer broadcast, offer acceptance/rejection/expiry, live tracking, QC photo, and QR proof of handoff.",
        "Bulk run workflow, including request, provider approval/rejection, pickup, distribution stop logging, served portion tracking, completion, and leftover stock return.",
        "Campaign and kitchen operations, including charity campaign creation, admin approval/rejection, volunteer staffing, kitchen preparation, meal batch tracking, distribution, and QR handoff.",
        "Profile and trust-score functions, including receiver/provider/shipper profile updates, trust penalties, restriction, ban, and token revocation behavior.",
        "Notification center functions, including realtime unread count, delivery offers, reservation updates, campaign updates, and read-state management.",
        "Admin and moderation functions for provider KYC review, campaign approval, user governance, violation handling, and operational monitoring.",
    ]:
        add_bullet(doc, item)

    add_para(doc, "Feature and Function not to be tested:", bold=True)
    for item in [
        "Large-scale production stress testing beyond the controlled capstone demo workload.",
        "Full third-party infrastructure reliability testing, such as external push-notification, geocoding, cloud storage, or hosting provider outages.",
        "Online payment settlement because FoodResQ focuses on donation, reservation, delivery, and distribution workflows.",
        "Formal biometric accuracy benchmarking; testing only verifies that the eKYC and face-enrollment gates behave correctly in the application workflow.",
        "Disaster recovery, multi-region failover, and production incident-response drills.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. Test Strategy", 2)
    add_heading(doc, "2.1 Testing Types", 3)
    testing_types = {
        "Functional Testing": [
            "Objective: Ensure FoodResQ functions comply with the requirement specifications and business rules.",
            "Technique: Black-box testing is used to evaluate the system from receiver, provider, shipper/volunteer, charity, and admin perspectives.",
            "Completion Criteria: All critical functional test cases must pass without any major defects.",
        ],
        "Integration Testing": [
            "Objective: Verify that frontend/mobile screens, NestJS APIs, PostgreSQL/PostGIS data, Redis locks, queues, and Socket.IO events work together correctly.",
            "Technique: Execute workflows that cross module boundaries, then check UI result, API response, database state, and notification behavior.",
            "Completion Criteria: No broken state transition or data mismatch remains in reservation, delivery, campaign, or bulk-run flows.",
        ],
        "User Acceptance Testing": [
            "Objective: Ensure the system supports realistic food-rescue operations for receivers, providers, shippers/volunteers, charities, and admins.",
            "Technique: Team members perform role-based demo scenarios using prepared accounts and seeded test data.",
            "Completion Criteria: All user stories and acceptance criteria required for the capstone demo are met.",
        ],
        "Security Testing": [
            "Objective: Validate authentication, authorization, banned/restricted account behavior, upload validation, and sensitive route protection.",
            "Technique: Attempt invalid credentials, expired tokens, wrong-role actions, unauthenticated access, and invalid file uploads.",
            "Completion Criteria: Unauthorized requests are rejected and protected operations require a valid session and correct role.",
        ],
        "Performance Testing": [
            "Objective: Ensure high-risk flows respond within acceptable demo limits.",
            "Technique: Measure nearby listing search, reservation creation, delivery offer broadcast, and realtime tracking under controlled test data.",
            "Completion Criteria: Main user-facing actions complete without noticeable blocking or inconsistent state.",
        ],
    }
    for name, bullets in testing_types.items():
        add_bullet(doc, name)
        for bullet in bullets:
            add_bullet(doc, bullet, level=1)

    add_heading(doc, "2.2 Test Levels", 3)
    make_table(
        doc,
        ["Type of Tests", "Unit", "Integration", "System", "Acceptance"],
        [
            ["Functionality", "X", "X", "X", "X"],
            ["User Interface", "", "", "X", "X"],
            ["Data Validation", "X", "X", "X", "X"],
            ["Role-based Access", "X", "X", "X", "X"],
            ["Realtime/Notification", "", "X", "X", "X"],
            ["Performance", "", "X", "X", ""],
        ],
        widths=[2.05, 1.0, 1.1, 1.05, 1.1],
    )
    add_caption(doc, "Table 1. Test levels table")

    add_heading(doc, "2.3 Supporting Tools", 3)
    make_table(
        doc,
        ["Purpose", "Tool", "Vendor/In-house", "Version"],
        [
            ["Test case management", "Microsoft Excel", "Microsoft", "Desktop / latest"],
            ["Manual test execution", "Chrome, Edge, Android Emulator, Expo Go", "Microsoft / Google / Expo", "Latest"],
            ["API verification", "Swagger / Postman", "In-house / Postman", "Latest"],
            ["Defect tracking", "GitHub Issues", "GitHub", "Cloud version"],
            ["Development environment", "Visual Studio Code", "Microsoft", "Latest"],
        ],
        widths=[1.8, 2.25, 1.45, 1.3],
    )
    add_caption(doc, "Table 2. Supporting tools table")

    add_heading(doc, "3. Test Plan", 2)
    add_heading(doc, "3.1 Human Resources", 3)
    make_table(
        doc,
        ["Worker/Doer", "Role", "Specific Responsibilities/Comments"],
        [
            ["QA Team", "Tester", "Prepare manual test cases, execute system and acceptance tests, record actual results"],
            ["Backend Developer", "Developer", "Support API, service rule, database, queue, and realtime issue investigation"],
            ["Web/Mobile Developer", "Developer", "Support UI, form validation, map, QR, and role-based screen issue investigation"],
            ["Project Lead", "Reviewer", "Review scope, milestone, defect severity, and final report consistency"],
        ],
        widths=[1.55, 1.25, 4.05],
    )
    add_caption(doc, "Table 3. Human resources table")

    add_heading(doc, "3.2 Test Environment", 3)
    make_table(
        doc,
        ["Purpose", "Tool", "Provider", "Version"],
        [
            ["Web application testing", "Chrome / Edge", "Google / Microsoft", "Latest"],
            ["Mobile application testing", "Expo Go / Android Emulator", "Expo / Google", "Latest"],
            ["Backend API", "NestJS API server", "In-house", "Latest project build"],
            ["Database", "PostgreSQL + PostGIS", "Open-source", "15.x+"],
            ["Realtime and background jobs", "Socket.IO, Redis, BullMQ", "Open-source", "Latest"],
            ["Source control", "Git", "Open-source", "Latest"],
        ],
        widths=[1.8, 2.25, 1.45, 1.3],
    )
    add_caption(doc, "Table 4. Test environment table")

    add_heading(doc, "3.3 Test Milestones", 3)
    make_table(
        doc,
        ["Milestone Task", "Start Date", "End Date"],
        [
            ["Test planning", "01/08/2026", "03/08/2026"],
            ["Test case development", "04/08/2026", "08/08/2026"],
            ["Manual system testing", "09/08/2026", "12/08/2026"],
            ["Acceptance testing", "13/08/2026", "14/08/2026"],
            ["Testing closure and reporting", "15/08/2026", "15/08/2026"],
        ],
        widths=[3.6, 1.55, 1.55],
    )
    add_caption(doc, "Table 5. Test milestones table")

    add_heading(doc, "4. Test Cases", 2)
    add_bullet(doc, "Manual System and Acceptance Test Cases: Refer to FoodResQ_Report5_Test_Report_v7.xlsx.")
    add_bullet(doc, "Detailed test case sheets cover Login, Register eKYC, Logout, Password, Profile Trust, Admin Controls, Food Listings, Reservations QR, Delivery Shipper, Campaign Kitchen Ops, and Bulk Notifications.")
    add_bullet(doc, "Unit test execution is maintained separately from this manual test report and should only be included if the team decides to submit automated evidence.")

    add_heading(doc, "5. Test Reports", 2)
    add_bullet(doc, "System Testing, Acceptance Testing")
    add_metadata_table(doc, "Test Statistics")

    stats_rows = []
    total = 0
    for idx, (code, module, count) in enumerate(MODULES, start=1):
        stats_rows.append([idx, f"{code} - {module}", count, 0, 0, 0, count])
        total += count
    stats_rows.append(["Sub total", "", total, 0, 0, 0, total])
    make_table(
        doc,
        ["No", "Module code", "Passed", "Failed", "Pending", "N/A", "Number of test cases"],
        stats_rows,
        widths=[0.55, 2.55, 0.65, 0.65, 0.75, 0.55, 1.1],
        blue_header=True,
        subtotal_row=True,
    )
    add_caption(doc, "Table 6. Test statistics report table")

    make_table(
        doc,
        ["Metric", "Value"],
        [
            ["Test coverage", "100.00%"],
            ["Test successful coverage", "100.00%"],
            ["Normal case", "43%"],
            ["Abnormal case", "41%"],
            ["Boundary case", "16%"],
        ],
        widths=[3.4, 3.2],
    )
    add_caption(doc, "Table 7. Test coverage summary")

    add_para(
        doc,
        "The result values above reflect the current manual test report workbook. During final manual execution, the tester should update Passed, Failed, Pending, and N/A values in both the workbook and this summary if any test result changes.",
        italic=True,
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "FoodResQ - Report 5 Test Documentation"
    for run in footer.runs:
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        run.font.size = Pt(9)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build_doc()
