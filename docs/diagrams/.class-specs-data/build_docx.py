"""
Build class-specs.docx by parsing 12 class-*.puml files and merging them with
the matching class-specs-data/*.json descriptions.
"""
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(r"d:\Do_An\foodresq\docs\diagrams")
DATA_DIR = BASE / ".class-specs-data"
OUT = BASE / "class-specs.docx"

# ----------- parse puml --------------------------------------------------

CLASS_BLOCK_RE = re.compile(
    r"class\s+(?P<name>\w+)\s*(?P<stereo><<[^>]+>>)?\s*\{(?P<body>.*?)\}",
    re.DOTALL,
)
METHOD_LINE_RE = re.compile(r"^\s*([+\-#~])\s*(.+?)\s*$")
SECTION_RE = re.compile(
    r"^'\u2500+\s*(?P<title>.+?)\s*\u2500+", re.UNICODE
)


def parse_puml(path: Path):
    text = path.read_text(encoding="utf-8")
    blocks = []
    current_title = None
    current_section_lines = None

    for line in text.splitlines():
        sm = SECTION_RE.match(line)
        if sm:
            if current_section_lines is not None and current_title:
                blocks.append((current_title, "\n".join(current_section_lines)))
            current_title = sm.group("title").strip()
            current_section_lines = None
            continue
        if line.strip().startswith("@startuml"):
            current_section_lines = [line]
            continue
        if line.strip().startswith("@enduml"):
            if current_section_lines is not None:
                current_section_lines.append(line)
                blocks.append((current_title or "(untitled)", "\n".join(current_section_lines)))
                current_section_lines = None
                current_title = None
            continue
        if current_section_lines is not None:
            current_section_lines.append(line)

    if current_section_lines is not None and current_title:
        blocks.append((current_title, "\n".join(current_section_lines)))

    out = []
    for title, body in blocks:
        classes = {}
        for m in CLASS_BLOCK_RE.finditer(body):
            cls = {"name": m.group("name"), "stereo": (m.group("stereo") or "").strip(),
                   "methods": [], "fields": []}
            for ln in m.group("body").splitlines():
                stripped = ln.strip()
                if not stripped:
                    continue
                mm = METHOD_LINE_RE.match(stripped)
                if mm:
                    cls["methods"].append(stripped)
                elif ":" in stripped and any(s in stripped for s in ["+", "-", "#", "~"]):
                    cls["fields"].append(stripped)
            classes[m.group("name")] = cls
        out.append((title, classes))
    return out


# ----------- helpers ----------------------------------------------------

def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip().lower()


def class_order_key(name):
    if name.endswith("Controller"):
        return 0
    if name.endswith("Service"):
        return 1
    if name.endswith("Cron"):
        return 2
    if name.endswith("Dto"):
        return 9
    return 5


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_styled_heading(doc, text, level, color="1F4E79"):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_table(doc, headers, rows, widths_in=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.autofit = False
    if widths_in:
        for i, w in enumerate(widths_in):
            for cell in t.columns[i].cells:
                cell.width = Inches(w)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        set_cell_bg(c, "1F4E79")
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = ""
            p = c.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Consolas" if ci == 0 else "Calibri"
            if ri % 2 == 0:
                set_cell_bg(c, "F2F2F2")
    return t


# ----------- main -------------------------------------------------------

def main():
    puml_files = sorted(BASE.glob("class-*.puml"))
    # Map puml name (e.g. class-admin.puml) to its json sibling (admin.json)
    data_files = {}
    for p in DATA_DIR.glob("*.json"):
        with p.open(encoding="utf-8") as f:
            d = json.load(f)
        if d:
            # Use the only top-level key as the mapping
            for k in d:
                data_files[k] = p
                break

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # Title page
    title = doc.add_paragraph()
    title.alignment = 1
    r = title.add_run("FoodResQ - Class Diagram Specifications")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor.from_string("1F4E79")

    sub = doc.add_paragraph()
    sub.alignment = 1
    sr = sub.add_run("Method-by-method description for 12 class diagrams\n"
                     "Capstone SP26SE088 - FPT University")
    sr.italic = True
    sr.font.size = Pt(12)

    doc.add_paragraph()
    intro = doc.add_paragraph()
    ir = intro.add_run(
        "This document is the textual companion to the 12 class diagrams in "
        "docs/diagrams/class-*.puml. Each entry lists every public and private "
        "method with its purpose, parameter handling, and notable side effects "
        "(transactions, trust-score deltas, socket emissions, file storage)."
    )
    ir.font.size = Pt(11)

    add_styled_heading(doc, "Files covered", level=1)
    for p in puml_files:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(p.name)
        run.font.name = "Consolas"
        run.font.size = Pt(10)

    doc.add_page_break()

    for puml in puml_files:
        sections = parse_puml(puml)
        json_path = data_files.get(puml.name)
        if not json_path:
            print(f"  WARN no json for {puml.name}")
            continue
        data = json.loads(json_path.read_text(encoding="utf-8"))
        # Drill down to inner: data[filename][section]
        inner = data.get(puml.name, {})

        # If a puml has only one startuml block and no section header,
        # fall back to the only key in inner.
        if len(sections) == 1 and sections[0][0] == "(untitled)" and len(inner) == 1:
            only_key = list(inner.keys())[0]
            sections = [(only_key, sections[0][1])]

        add_styled_heading(doc, puml.name, level=1, color="1F4E79")

        for title_text, classes in sections:
            data_block = None
            if isinstance(inner, dict):
                for k, v in inner.items():
                    if _norm(k) == _norm(title_text):
                        data_block = v
                        break
            if not isinstance(data_block, dict):
                data_block = {}

            add_styled_heading(doc, title_text, level=2, color="2E75B6")

            explanation = data_block.get("explanation") if isinstance(data_block, dict) else None
            if explanation:
                ep = doc.add_paragraph()
                er = ep.add_run("Overview: ")
                er.bold = True
                er.font.size = Pt(11)
                er.font.color.rgb = RGBColor.from_string("404040")
                er2 = ep.add_run(explanation)
                er2.font.size = Pt(11)
                er2.italic = True
                er2.font.color.rgb = RGBColor.from_string("404040")

            classes_data = data_block.get("classes", {}) if isinstance(data_block, dict) else {}

            sorted_names = sorted(classes.keys(), key=class_order_key)

            for cls_name in sorted_names:
                cls = classes[cls_name]
                meta = classes_data.get(cls_name, {}) if isinstance(classes_data, dict) else {}
                cls_desc = meta.get("description", "") if isinstance(meta, dict) else ""
                cls_methods = meta.get("methods", {}) if isinstance(meta, dict) else {}

                add_styled_heading(doc, cls_name, level=3, color="5B9BD5")

                if cls_desc:
                    add_para(doc, cls_desc, italic=True, color="404040", size=10)

                # PrismaService exposes only model delegates — surface them as a
                # compact list of bullets rather than a sparse table.
                if cls_name == "PrismaService":
                    if cls["methods"]:
                        add_para(doc, "Model delegates exposed by PrismaClient:", bold=True, size=10, color="595959")
                        for f in cls["methods"]:
                            add_para(doc, "  - " + f, size=9, color="595959")
                    continue

                if cls["methods"]:
                    rows = []
                    for sig in cls["methods"]:
                        desc = cls_methods.get(sig, "") if isinstance(cls_methods, dict) else ""
                        rows.append([sig, desc])
                    add_table(doc, ["Method", "Description"], rows, widths_in=[3.5, 4.0])

                if cls["fields"]:
                    add_para(doc, "Fields:", bold=True, size=10, color="595959")
                    for f in cls["fields"]:
                        add_para(doc, f, size=9, color="595959")

        doc.add_paragraph()

    tmp = OUT.with_name(f"_new_{OUT.name}")
    doc.save(tmp)
    if OUT.exists():
        try:
            OUT.unlink()
            tmp.rename(OUT)
        except PermissionError:
            # Word may be holding OUT. The new file is at tmp; user can rename.
            print(f"WARN: {OUT.name} is locked; wrote {tmp.name} instead.")
            return
    else:
        tmp.rename(OUT)
    print(f"Wrote {OUT}")
    print(f"Total files: {len(puml_files)}")


if __name__ == "__main__":
    main()
